#!/usr/bin/env python3
"""
Logistics Document Generator

This script generates multi-page shipping placards from Excel data using Word templates.
Reads from Data folder, uses Template folder for templates, outputs to Placards folder.

Security Features:
- Input validation and sanitization
- Path traversal protection
- File size and type validation
- Rate limiting and resource management
- Comprehensive error handling and logging
"""

import os
import sys
import csv
import re
import logging
import threading
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple, Any, cast

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Import security utilities
from .security import (
    InputValidator, PathSanitizer, SecureFileHandler, SecurityConfig,
    SecurityError, RateLimiter, security_logger
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('application.log'),
        logging.StreamHandler()
    ]
)
app_logger = logging.getLogger('placard_generator')


class PlacardGenerator:
    """High-performance logistics document generator with comprehensive security"""
    
    def __init__(self):
        self.df: Optional[pd.DataFrame] = None
        self.required_columns = [
            'Shipment Nbr', 'DO #', 'Label Type', 'Order Type', 
            'Pmt Term', 'Start Ship', 'VAS', 'Ship To', 'PO', 'Original Qty'
        ]
        self.data_folder = "Data"
        self.template_folder = "Template"
        self.output_folder = "Placards"
        self.log_folder = "Logs"
        self.log_file = None
        
        # Security components
        self.data_handler = SecureFileHandler(self.data_folder)
        self.template_handler = SecureFileHandler(self.template_folder)
        self.output_handler = SecureFileHandler(self.output_folder)
        self.log_handler = SecureFileHandler(self.log_folder)
        self.rate_limiter = RateLimiter(SecurityConfig.MAX_OPERATIONS_PER_MINUTE)
        
        # Thread safety
        self._processing_lock = threading.Lock()
        self._is_processing = False
        
        # Performance monitoring
        self._start_time = None
        self._processed_count = 0
        
    def get_timestamp(self) -> str:
        """Get a readable timestamp for console output"""
        return datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
    
    def print_with_timestamp(self, message: str) -> None:
        """Print message with timestamp prefix"""
        print(f"{self.get_timestamp()} {message}")
        
    def setup_directories(self) -> bool:
        """Ensure required directories exist"""
        try:
            for folder in [self.data_folder, self.template_folder, self.output_folder, self.log_folder]:
                if not os.path.exists(folder):
                    os.makedirs(folder)
                    self.print_with_timestamp(f"Created directory: {folder}")
            return True
        except Exception as e:
            self.print_with_timestamp(f"Error creating directories: {e}")
            return False
    
    def initialize_log(self) -> bool:
        """Initialize CSV log file with headers"""
        try:
            # Create log filename with timestamp at front - more readable format
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            self.log_file = os.path.join(self.log_folder, f"{timestamp}-placard_processing_log.csv")
            
            # Create CSV with headers
            with open(self.log_file, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow([
                    'Timestamp',
                    'Session_ID',
                    'Event_Type',
                    'Shipment_Number',
                    'DO_Count',
                    'Records_Found',
                    'Status',
                    'Output_File',
                    'Error_Message',
                    'Processing_Mode',
                    'Duration_Seconds'
                ])
            
            self.print_with_timestamp(f"Logging to: {self.log_file}")
            return True
            
        except Exception as e:
            self.print_with_timestamp(f"Warning: Could not initialize log file: {e}")
            return False
    
    def log_event(self, event_type: str, shipment_number: Optional[str] = None, 
                  do_count: Optional[int] = None, records_found: Optional[int] = None,
                  status: str = "SUCCESS", output_file: Optional[str] = None,
                  error_message: Optional[str] = None, processing_mode: Optional[str] = None,
                  duration: Optional[float] = None) -> None:
        """Log an event to the CSV file"""
        if not self.log_file:
            return
            
        try:
            with open(self.log_file, 'a', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow([
                    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    datetime.now().strftime("%Y%m%d_%H%M%S"),  # Session ID based on startup time
                    event_type,
                    shipment_number or "",
                    do_count or "",
                    records_found or "",
                    status,
                    output_file or "",
                    error_message or "",
                    processing_mode or "",
                    f"{duration:.2f}" if duration else ""
                ])
        except Exception as e:
            self.print_with_timestamp(f"Warning: Could not write to log file: {e}")
    
    def find_excel_file(self) -> Optional[str]:
        """Find Excel file starting with 'WM-SPN-CUS105 Open Order Report' in Data folder"""
        try:
            # Securely list files in data directory
            xlsx_files = self.data_handler.safe_list_files("WM-SPN-CUS105 Open Order Report*.xlsx")
            xls_files = self.data_handler.safe_list_files("WM-SPN-CUS105 Open Order Report*.xls")
            
            all_files = xlsx_files + xls_files
            
            if not all_files:
                error_msg = f"ERROR: No Excel file found in '{self.data_folder}' folder starting with 'WM-SPN-CUS105 Open Order Report'"
                self.print_with_timestamp(error_msg)
                security_logger.warning(f"Excel file not found in {self.data_folder}")
                return None
            
            # Additional security validation
            valid_files = []
            for file_path in all_files:
                # Validate file extension
                if not SecurityConfig.validate_file_extension(file_path, SecurityConfig.ALLOWED_EXCEL_EXTENSIONS):
                    security_logger.warning(f"Invalid file extension: {file_path}")
                    continue
                    
                # Validate file size
                if not SecurityConfig.validate_file_size(file_path, SecurityConfig.MAX_EXCEL_FILE_SIZE):
                    security_logger.warning(f"File too large: {file_path}")
                    continue
                    
                # Calculate and log file hash for integrity
                file_hash = self.data_handler.calculate_file_hash(file_path)
                if file_hash:
                    app_logger.info(f"Processing file: {file_path} (SHA256: {file_hash[:16]}...)")
                    
                valid_files.append(file_path)
            
            if not valid_files:
                self.print_with_timestamp("ERROR: No valid Excel files found after security validation")
                return None
            
            if len(valid_files) > 1:
                self.print_with_timestamp(f"Multiple Excel files found. Using: {valid_files[0]}")
                security_logger.info(f"Multiple files available, selected: {valid_files[0]}")
            
            return valid_files[0]
            
        except SecurityError as e:
            security_logger.error(f"Security error in file discovery: {e}")
            self.print_with_timestamp(f"Security error: {e}")
            return None
        except Exception as e:
            app_logger.error(f"Unexpected error in file discovery: {e}")
            self.print_with_timestamp(f"Unexpected error: {e}")
            return None
    
    def validate_do_number(self, do_num: Any) -> bool:
        """Validate DO # format using secure validation"""
        try:
            return InputValidator.validate_do_number(do_num)
        except Exception as e:
            security_logger.error(f"DO number validation error: {e}")
            return False
    
    def validate_shipment_number(self, shipment_num: Any) -> bool:
        """Validate shipment number using secure validation"""
        try:
            return InputValidator.validate_shipment_number(shipment_num)
        except Exception as e:
            security_logger.error(f"Shipment number validation error: {e}")
            return False
    
    def validate_data_integrity(self, df: pd.DataFrame) -> bool:
        """Comprehensive data validation with security checks"""
        try:
            validation_errors = []
            
            # Check required columns
            missing_columns = [col for col in self.required_columns if col not in df.columns]
            if missing_columns:
                validation_errors.append(f"Missing required columns: {missing_columns}")
            
            # Validate data types and content
            for index, row in df.iterrows():
                # Validate shipment number
                if not self.validate_shipment_number(row.get('Shipment Nbr')):
                    validation_errors.append(f"Invalid shipment number at row {index}")
                
                # Validate DO number
                if not self.validate_do_number(row.get('DO #')):
                    validation_errors.append(f"Invalid DO number at row {index}")
                
                # Validate text fields
                text_fields = ['Ship To', 'PO', 'Label Type', 'Order Type', 'Pmt Term']
                for field in text_fields:
                    if not InputValidator.validate_text_field(row.get(field), max_length=500):
                        validation_errors.append(f"Invalid {field} at row {index}")
                
                # Validate numeric fields
                if not InputValidator.validate_numeric_field(row.get('Original Qty')):
                    validation_errors.append(f"Invalid quantity at row {index}")
            
            # Log validation results
            if validation_errors:
                for error in validation_errors[:10]:  # Log first 10 errors
                    security_logger.warning(f"Data validation error: {error}")
                if len(validation_errors) > 10:
                    security_logger.warning(f"... and {len(validation_errors) - 10} more validation errors")
                return False
            
            app_logger.info(f"Data validation passed for {len(df)} records")
            return True
            
        except Exception as e:
            security_logger.error(f"Data validation failed with exception: {e}")
            return False
    
    def load_and_prepare_data(self) -> bool:
        """Load Excel file and prepare data with comprehensive security validation"""
        with self._processing_lock:
            if self._is_processing:
                self.print_with_timestamp("ERROR: Another data loading operation is in progress")
                return False
            self._is_processing = True
        
        try:
            self.print_with_timestamp("Loading and preparing data...")
            start_time = datetime.now()
            self._start_time = start_time
            
            # Rate limiting check
            if not self.rate_limiter.allow_operation():
                error_msg = "Rate limit exceeded. Please wait before retrying."
                self.print_with_timestamp(f"ERROR: {error_msg}")
                security_logger.warning("Data loading rate limit exceeded")
                self.log_event("DATA_LOAD", status="FAILED", error_message=error_msg)
                return False
            
            # Find Excel file with security validation
            excel_file = self.find_excel_file()
            if not excel_file:
                self.log_event("DATA_LOAD", status="FAILED", error_message="Excel file not found")
                return False
            
            # Additional file validation before processing
            if not self.data_handler.safe_file_exists(excel_file):
                error_msg = "Excel file validation failed"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                security_logger.error(f"File validation failed: {excel_file}")
                self.log_event("DATA_LOAD", status="FAILED", error_message=error_msg)
                return False
                
            # Load Excel file with error handling
            self.print_with_timestamp(f"Loading file: {excel_file}")
            try:
                df = pd.read_excel(excel_file, engine='openpyxl')
            except Exception as read_error:
                # Try with xlrd engine for .xls files
                try:
                    df = pd.read_excel(excel_file, engine='xlrd')
                except Exception:
                    error_msg = f"Failed to read Excel file with both engines: {read_error}"
                    self.print_with_timestamp(f"ERROR: {error_msg}")
                    security_logger.error(f"Excel read error: {read_error}")
                    self.log_event("DATA_LOAD", status="FAILED", error_message=str(read_error))
                    return False
            
            initial_count = len(df)
            self.print_with_timestamp(f"Loaded {initial_count} rows from Excel file")
            
            # Check for reasonable data size
            if initial_count > SecurityConfig.MAX_RECORDS_PER_BATCH:
                error_msg = f"Dataset too large ({initial_count} records). Maximum allowed: {SecurityConfig.MAX_RECORDS_PER_BATCH}"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                security_logger.warning(f"Large dataset detected: {initial_count} records")
                self.log_event("DATA_LOAD", status="FAILED", error_message=error_msg)
                return False
            
            # Comprehensive data validation
            if not self.validate_data_integrity(df):
                error_msg = "Data integrity validation failed"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                self.log_event("DATA_LOAD", status="FAILED", error_message=error_msg)
                return False
            
            # Filter out rows with empty Shipment Nbr
            df = df[df['Shipment Nbr'].notna()]
            shipment_series = cast(pd.Series, df['Shipment Nbr'].astype(str))
            df = df[shipment_series.str.strip() != '']
            empty_removed = initial_count - len(df)
            
            # Validate DO # format using secure validation
            before_do_filter = len(df)
            do_series = cast(pd.Series, df['DO #'])
            df = df[do_series.apply(self.validate_do_number)]
            invalid_do_removed = before_do_filter - len(df)
            
            # Validate shipment numbers
            before_shipment_filter = len(df)
            shipment_series = cast(pd.Series, df['Shipment Nbr'])
            df = df[shipment_series.apply(self.validate_shipment_number)]
            invalid_shipment_removed = before_shipment_filter - len(df)
            
            # Assign to instance variable
            self.df = cast(pd.DataFrame, df)
            
            # Calculate processing time
            duration = (datetime.now() - start_time).total_seconds()
            
            if self.df is not None:
                final_count = len(self.df)
                self.print_with_timestamp(f"Data preparation summary:")
                self.print_with_timestamp(f"  - Initial records: {initial_count}")
                self.print_with_timestamp(f"  - Removed empty: {empty_removed}")
                self.print_with_timestamp(f"  - Removed invalid DO#: {invalid_do_removed}")
                self.print_with_timestamp(f"  - Removed invalid shipments: {invalid_shipment_removed}")
                self.print_with_timestamp(f"  - Final dataset: {final_count} rows ready for processing")
                
                # Log successful data load
                self.log_event(
                    "DATA_LOAD", 
                    records_found=final_count,
                    status="SUCCESS",
                    error_message=f"Processed {initial_count} -> {final_count} records",
                    duration=duration
                )
                
                app_logger.info(f"Data loaded successfully: {final_count} valid records")
            return True
            
        except Exception as e:
            duration = (datetime.now() - start_time).total_seconds() if start_time else 0
            error_msg = f"ERROR loading Excel file: {e}"
            self.print_with_timestamp(error_msg)
            app_logger.error(f"Data loading failed: {e}", exc_info=True)
            self.log_event("DATA_LOAD", status="FAILED", error_message=str(e), duration=duration)
            return False
            
        finally:
            with self._processing_lock:
                self._is_processing = False
    
    def format_date(self, date_value: Any) -> str:
        """Format date as MM/DD/YYYY"""
        if pd.isna(date_value):
            return ""
        
        try:
            if isinstance(date_value, str):
                # Try to parse string date
                date_obj = pd.to_datetime(date_value)
            else:
                date_obj = pd.to_datetime(date_value)
            return date_obj.strftime("%m/%d/%Y")
        except:
            return str(date_value)
    
    def get_vas_value(self, vas_raw: Any) -> str:
        """Convert VAS value to 'VAS' or 'NOT VAS'"""
        if pd.isna(vas_raw):
            return "NOT VAS"
        vas_str = str(vas_raw).strip().upper()
        return "VAS" if vas_str == "Y" else "NOT VAS"
    
    def replace_placeholders_in_document(self, doc: Any, replacements: Dict[str, str]) -> None:
        """Replace all placeholders in the document while preserving formatting"""
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    self.replace_placeholder_in_paragraph(paragraph, placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self.replace_placeholder_in_paragraph(paragraph, placeholder, value)
        
        # Replace in headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self.replace_placeholder_in_paragraph(paragraph, placeholder, value)
        
        # Replace in footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self.replace_placeholder_in_paragraph(paragraph, placeholder, value)
    
    def replace_placeholder_in_paragraph(self, paragraph: Any, placeholder: str, replacement: str) -> None:
        """Replace placeholder while preserving mixed formatting"""
        if placeholder not in paragraph.text:
            return
        
        # Try simple replacement first (best case - placeholder is in one run)
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
                return
        
        # Complex case: placeholder spans multiple runs
        # We need to reconstruct the paragraph preserving different formatting for different parts
        self.replace_across_runs(paragraph, placeholder, replacement)
    
    def replace_across_runs(self, paragraph: Any, placeholder: str, replacement: str) -> None:
        """Replace placeholder that spans across multiple runs while preserving original formatting"""
        
        # Store all original runs with their formatting
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None
            })
        
        # Get the full text and find placeholder position
        full_text = paragraph.text
        placeholder_start = full_text.find(placeholder)
        if placeholder_start == -1:
            return
        
        placeholder_end = placeholder_start + len(placeholder)
        
        # Split text into three parts: before placeholder, placeholder, after placeholder
        before_text = full_text[:placeholder_start]
        after_text = full_text[placeholder_end:]
        
        # Find the run that contains most of the placeholder for formatting reference
        char_pos = 0
        placeholder_run_format = None
        
        for run_info in original_runs:
            run_end = char_pos + len(run_info['text'])
            if char_pos <= placeholder_start < run_end:
                placeholder_run_format = run_info
                break
            char_pos = run_end
        
        # Fallback to first run if we can't find the placeholder run
        if placeholder_run_format is None and original_runs:
            placeholder_run_format = original_runs[0]
        
        # Clear the paragraph
        paragraph.clear()
        
        # Method 1: Try to reconstruct the original structure intelligently
        if before_text or after_text:
            self.reconstruct_mixed_formatting(paragraph, before_text, replacement, after_text, original_runs, placeholder_run_format)
        else:
            # Simple case: just the placeholder
            new_run = paragraph.add_run(replacement)
            if placeholder_run_format:
                self.apply_run_formatting(new_run, placeholder_run_format)
    
    def reconstruct_mixed_formatting(self, paragraph: Any, before_text: str, replacement: str, after_text: str, original_runs: List[Dict[str, Any]], placeholder_format: Optional[Dict[str, Any]]) -> None:
        """Reconstruct paragraph with mixed formatting preserved"""
        
        # For now, use a simplified approach that preserves the most important formatting
        # Add before text (try to preserve original formatting)
        if before_text:
            # Use the first run's formatting for the before text
            before_run = paragraph.add_run(before_text)
            if original_runs:
                self.apply_run_formatting(before_run, original_runs[0])
        
        # Add replacement text with placeholder formatting  
        replacement_run = paragraph.add_run(replacement)
        if placeholder_format:
            self.apply_run_formatting(replacement_run, placeholder_format)
        
        # Add after text (try to preserve original formatting)
        if after_text:
            # Use the last run's formatting for the after text
            after_run = paragraph.add_run(after_text)
            if original_runs:
                self.apply_run_formatting(after_run, original_runs[-1])
    
    def apply_run_formatting(self, run: Any, format_info: Dict[str, Any]) -> None:
        """Apply formatting from format_info dictionary to a run"""
        if format_info.get('bold') is not None:
            run.bold = format_info['bold']
        if format_info.get('italic') is not None:
            run.italic = format_info['italic']
        if format_info.get('underline') is not None:
            run.underline = format_info['underline']
        if format_info.get('font_name'):
            run.font.name = format_info['font_name']
        if format_info.get('font_size'):
            run.font.size = format_info['font_size']
        if format_info.get('font_color'):
            run.font.color.rgb = format_info['font_color']
    
    def replace_in_paragraph(self, paragraph: Any, placeholder: str, replacement: str) -> None:
        """Replace placeholder in paragraph while preserving formatting"""
        if placeholder not in paragraph.text:
            return
        
        # Method 1: Try to replace within existing runs first (best formatting preservation)
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
                return
        
        # Method 2: Handle placeholder spanning multiple runs with smart formatting detection
        full_text = paragraph.text
        if placeholder not in full_text:
            return
        
        # Find placeholder position
        placeholder_start = full_text.find(placeholder)
        placeholder_end = placeholder_start + len(placeholder)
        
        # Collect all runs and their positions
        runs_info = []
        current_pos = 0
        
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            runs_info.append({
                'run': run,
                'start': run_start,
                'end': run_end,
                'text': run.text
            })
            current_pos = run_end
        
        # Find which runs contain the placeholder
        affected_runs = []
        for run_info in runs_info:
            if (run_info['start'] < placeholder_end and run_info['end'] > placeholder_start):
                affected_runs.append(run_info)
        
        if not affected_runs:
            return
        
        # Use the formatting from the run that contains the start of the placeholder
        default_run = None
        for run_info in affected_runs:
            if run_info['start'] <= placeholder_start < run_info['end']:
                default_run = run_info['run']
                break
        
        if default_run is None:
            default_run = affected_runs[0]['run']
        
        # Replace with mixed formatting preservation
        self.replace_with_mixed_formatting(paragraph, placeholder, replacement, default_run)
    
    def apply_formatting(self, run: Any, bold: Optional[bool], italic: Optional[bool], underline: Optional[bool], font_name: Optional[str], font_size: Optional[Any], font_color: Optional[Any]) -> None:
        """Apply formatting to a run"""
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_color:
            run.font.color.rgb = font_color
    
    def replace_with_mixed_formatting(self, paragraph: Any, placeholder: str, replacement: str, default_run: Any) -> None:
        """Replace placeholder while preserving mixed formatting in the paragraph"""
        # Get the full paragraph text
        full_text = paragraph.text
        
        # Find placeholder position
        placeholder_start = full_text.find(placeholder)
        if placeholder_start == -1:
            return
        
        placeholder_end = placeholder_start + len(placeholder)
        
        # Split into before, placeholder, and after
        before_text = full_text[:placeholder_start]
        after_text = full_text[placeholder_end:]
        
        # Clear paragraph and rebuild with formatting
        paragraph.clear()
        
        # Add before text (preserve original formatting from first run)
        if before_text:
            before_run = paragraph.add_run(before_text)
            # Copy formatting from first original run
            if paragraph.runs:  # This will be empty since we just cleared
                pass  # We'll use default formatting
        
        # Add replacement text with default run formatting
        replacement_run = paragraph.add_run(replacement)
        self.apply_formatting(
            replacement_run,
            default_run.bold if default_run else None,
            default_run.italic if default_run else None,
            default_run.underline if default_run else None,
            default_run.font.name if default_run else None,
            default_run.font.size if default_run else None,
            default_run.font.color.rgb if default_run and default_run.font.color.rgb else None
        )
    
    def copy_template_content(self, template_path: str) -> Optional[Any]:
        """Load template and return a copy"""
        try:
            if not os.path.exists(template_path):
                self.print_with_timestamp(f"ERROR: Template file not found: {template_path}")
                return None
            
            doc = Document(template_path)  # type: ignore
            return doc
        except Exception as e:
            self.print_with_timestamp(f"ERROR loading template: {e}")
            return None
    
    def copy_formatted_content(self, source_doc: Any, target_doc: Any) -> None:
        """Copy all content from source document to target document while preserving formatting"""
        try:
            # Copy paragraphs with full formatting
            for paragraph in source_doc.paragraphs:
                new_paragraph = target_doc.add_paragraph()
                
                # Copy paragraph-level formatting
                new_paragraph.alignment = paragraph.alignment
                new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
                new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
                new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
                
                # Copy all runs with their formatting
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    
                    # Copy run formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
            
            # Copy tables with formatting
            for table in source_doc.tables:
                new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                
                # Copy table style
                if table.style:
                    new_table.style = table.style
                
                # Copy cell content and formatting
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_cell = new_table.cell(i, j)
                        
                        # Clear default paragraph and copy all paragraphs from source cell
                        new_cell.paragraphs[0].clear()
                        
                        for paragraph in cell.paragraphs:
                            if paragraph == cell.paragraphs[0]:
                                # Use the existing first paragraph
                                target_paragraph = new_cell.paragraphs[0]
                            else:
                                # Add new paragraph for additional ones
                                target_paragraph = new_cell.add_paragraph()
                            
                            # Copy paragraph formatting
                            target_paragraph.alignment = paragraph.alignment
                            
                            # Copy runs with formatting
                            for run in paragraph.runs:
                                new_run = target_paragraph.add_run(run.text)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                if run.font.name:
                                    new_run.font.name = run.font.name
                                if run.font.size:
                                    new_run.font.size = run.font.size
                                if run.font.color.rgb:
                                    new_run.font.color.rgb = run.font.color.rgb
                                    
        except Exception as e:
            self.print_with_timestamp(f"Warning: Error copying formatted content: {e}")
            # Fallback to simple text copy
            for paragraph in source_doc.paragraphs:
                target_doc.add_paragraph(paragraph.text)
            for table in source_doc.tables:
                new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
    
    def _process_shipment_data(self, matching_records: pd.DataFrame) -> Optional[Dict[str, List[str]]]:
        """Process shipment data with validation and security checks"""
        try:
            # Group by DO # for processing
            do_groups = matching_records.groupby('DO #')
            do_data = {
                'do_numbers': [],
                'pos': [],
                'ship_to_info': [],
                'start_ship_dates': [],
                'vas_values': [],
                'quantities': [],
                'label_types': [],
                'order_types': [],
                'pmt_terms': []
            }
            
            for do_num, group in do_groups:
                # Validate DO number
                if not self.validate_do_number(do_num):
                    security_logger.warning(f"Invalid DO number in data: {do_num}")
                    continue
                    
                do_data['do_numbers'].append(str(int(float(str(do_num)))))
                
                # Collect unique PO numbers for this DO with validation
                po_list = group['PO'].dropna().unique().tolist()
                validated_pos = []
                for po in po_list:
                    if InputValidator.validate_text_field(po, max_length=100):
                        validated_pos.append(str(po).strip())
                    else:
                        security_logger.warning(f"Invalid PO number filtered: {po}")
                
                po_str = '\n'.join(validated_pos)
                do_data['pos'].append(po_str)
                
                # Get and validate other info from first record of group
                first_record = group.iloc[0]
                
                # Validate and sanitize text fields
                ship_to = str(first_record['Ship To']) if pd.notna(first_record['Ship To']) else ''
                if not InputValidator.validate_text_field(ship_to, max_length=500):
                    security_logger.warning(f"Invalid ship_to field: {ship_to[:50]}...")
                    ship_to = "INVALID_ADDRESS"
                
                label_type = str(first_record['Label Type']) if pd.notna(first_record['Label Type']) else ''
                if not InputValidator.validate_text_field(label_type, max_length=100):
                    security_logger.warning(f"Invalid label_type: {label_type}")
                    label_type = "STANDARD"
                
                order_type = str(first_record['Order Type']) if pd.notna(first_record['Order Type']) else ''
                if not InputValidator.validate_text_field(order_type, max_length=100):
                    security_logger.warning(f"Invalid order_type: {order_type}")
                    order_type = "STANDARD"
                
                pmt_term = str(first_record['Pmt Term']) if pd.notna(first_record['Pmt Term']) else ''
                if not InputValidator.validate_text_field(pmt_term, max_length=100):
                    security_logger.warning(f"Invalid pmt_term: {pmt_term}")
                    pmt_term = "NET30"
                
                # Validate quantity
                total_qty = group['Original Qty'].sum()
                if not InputValidator.validate_numeric_field(total_qty, min_val=0, max_val=1e6):
                    security_logger.warning(f"Invalid quantity: {total_qty}")
                    total_qty = 0
                
                do_data['ship_to_info'].append(ship_to)
                do_data['start_ship_dates'].append(self.format_date(first_record['Start Ship']))
                do_data['vas_values'].append(self.get_vas_value(first_record['VAS']))
                do_data['quantities'].append(f"{int(total_qty)} Units")
                do_data['label_types'].append(label_type)
                do_data['order_types'].append(order_type)
                do_data['pmt_terms'].append(pmt_term)
            
            if not do_data['do_numbers']:
                security_logger.error("No valid DO numbers found after validation")
                return None
                
            return do_data
            
        except Exception as e:
            app_logger.error(f"Error processing shipment data: {e}", exc_info=True)
            return None
    
    def _create_secure_document(self, template_path: str, shipment_num: str, do_data: Dict[str, List[str]]) -> Optional[Any]:
        """Create document with security validation"""
        try:
            # Load template securely
            doc = self.copy_template_content(template_path)
            if doc is None:
                return None
            
            # Process each DO as a separate page
            for i, (do_num, po_str, ship_to, start_ship, vas, qty, label_type, order_type, pmt_term) in enumerate(
                zip(do_data['do_numbers'], do_data['pos'], do_data['ship_to_info'], 
                    do_data['start_ship_dates'], do_data['vas_values'], do_data['quantities'], 
                    do_data['label_types'], do_data['order_types'], do_data['pmt_terms'])):
                
                # Add page break for subsequent pages
                if i > 0:
                    doc.add_page_break()
                    self.copy_formatted_content(doc, doc)
                
                # Prepare replacements with additional validation
                replacements = {
                    '{{Ship To}}': ship_to[:500],  # Limit length
                    '{{Shipment Nbr}}': shipment_num,
                    '{{PO}}': po_str[:1000],  # Limit length
                    '{{DO #}}': do_num,
                    '{{VAS}}': vas,
                    '{{Original Qty}}': qty,
                    '{{Label Type}}': label_type[:100],  # Limit length
                    '{{Order Type}}': order_type[:100],  # Limit length
                    '{{Pmt Term}}': pmt_term[:100],  # Limit length
                    '{{Start Ship}}': start_ship
                }
                
                # Replace placeholders
                self.replace_placeholders_in_document(doc, replacements)
            
            return doc
            
        except Exception as e:
            app_logger.error(f"Error creating document: {e}", exc_info=True)
            return None
    
    def _save_document_securely(self, doc: Any, shipment_num: str, do_count: int, 
                               record_count: int, start_time: datetime) -> bool:
        """Save document with security validation"""
        try:
            # Generate secure filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = f"Placard_{shipment_num}_{timestamp}"
            sanitized_filename = PathSanitizer.sanitize_filename(base_filename) + ".docx"
            
            # Use secure path joining
            output_path = PathSanitizer.safe_join_path(self.output_folder, sanitized_filename)
            
            # Ensure we can write to the output directory
            if not os.access(self.output_folder, os.W_OK):
                error_msg = f"No write permission to output directory: {self.output_folder}"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                security_logger.error(error_msg)
                return False
            
            # Save document with error handling
            doc.save(output_path)
            
            # Verify file was created and has reasonable size
            if not os.path.exists(output_path):
                error_msg = "Document save failed - file not created"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                return False
            
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                error_msg = "Document save failed - empty file created"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                os.remove(output_path)  # Clean up empty file
                return False
            
            # Calculate and log file hash for integrity
            file_hash = self.output_handler.calculate_file_hash(output_path)
            
            self.print_with_timestamp(f"✓ Successfully created: {sanitized_filename} ({file_size} bytes)")
            
            # Log successful processing
            duration = (datetime.now() - start_time).total_seconds()
            self.log_event("SHIPMENT_PROCESS", 
                          shipment_number=shipment_num,
                          do_count=do_count,
                          records_found=record_count,
                          status="SUCCESS",
                          output_file=sanitized_filename,
                          processing_mode="MANUAL",
                          duration=duration)
            
            if file_hash:
                app_logger.info(f"Document created: {sanitized_filename} (SHA256: {file_hash[:16]}...)")
            
            return True
            
        except SecurityError as e:
            error_msg = f"Security error saving document: {e}"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            security_logger.error(error_msg)
            return False
            
        except Exception as e:
            error_msg = f"Failed to save document: {e}"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            app_logger.error(f"Document save error: {e}", exc_info=True)
            return False
    
    def process_shipment(self, shipment_num: str) -> bool:
        """Process a single shipment number and generate placard"""
        self.print_with_timestamp(f"\nProcessing shipment: {shipment_num}")
        start_time = datetime.now()
        
        # Validate shipment number
        if not self.validate_shipment_number(shipment_num):
            error_msg = f"Invalid shipment number format: {shipment_num} (must be exactly 10 digits)"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num, 
                          status="FAILED", error_message=error_msg)
            return False
        
        # Ensure df is not None
        if self.df is None:
            error_msg = "No data loaded"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num, 
                          status="FAILED", error_message=error_msg)
            return False
        
        # Find shipment data - handle float values like 9010157586.0
        # Convert both the column and search value to integers for comparison
        df_shipment_clean = self.df['Shipment Nbr'].astype(float).astype(int).astype(str)
        shipment_data = self.df[df_shipment_clean == shipment_num]
        if shipment_data.empty:
            error_msg = f"No data found for shipment number: {shipment_num}"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num,
                          status="FAILED", error_message=error_msg)
            return False
        
        records_found = len(shipment_data)
        self.print_with_timestamp(f"Found {records_found} records for shipment {shipment_num}")
        
        # Get template path
        template_path = os.path.join(self.template_folder, "placard_template.docx")
        
        # Create main document by copying template
        main_doc = self.copy_template_content(template_path)
        if not main_doc:
            error_msg = "Failed to load template"
            self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num,
                          records_found=records_found, status="FAILED", error_message=error_msg)
            return False
        
        # Get shipment-level data (same for all pages)
        first_row = shipment_data.iloc[0]
        shipment_level_data = {
            'Shipment Nbr': str(first_row['Shipment Nbr']),
            'Label Type': str(first_row['Label Type']),
            'Order Type': str(first_row['Order Type']),
            'Pmt Term': str(first_row['Pmt Term']),
            'Start Ship': self.format_date(first_row['Start Ship']),
            'VAS': self.get_vas_value(first_row['VAS'])
        }
        
        # Group by DO #
        do_groups = shipment_data.groupby('DO #')
        total_dos = len(do_groups)
        
        self.print_with_timestamp(f"Processing {total_dos} DO #s for shipment {shipment_num}")
        
        # Process each DO # and create pages
        main_doc = None
        
        for do_index, (do_num, do_group) in enumerate(do_groups, 1):
            self.print_with_timestamp(f"  Processing DO # {do_num} ({do_index}/{total_dos})")
            
            # Get page-level data
            first_do_row = do_group.iloc[0]
            
            # Aggregate POs for this DO #
            # Cast to Series to access .dropna method
            po_series = cast(pd.Series, do_group['PO'])
            unique_pos = po_series.dropna().unique()
            po_list = '\n'.join([str(po) for po in unique_pos if str(po).strip()])
            
            # Calculate total Original Qty for this DO #
            total_original_qty = do_group['Original Qty'].sum()
            
            page_level_data = {
                'DO #': str(do_num),
                'Ship To': str(first_do_row['Ship To']),
                'PO': po_list,
                'Original Qty': str(int(total_original_qty)) if not pd.isna(total_original_qty) else '0'
            }
            
            # Combine all replacement data
            replacements = {
                '{{Ship To}}': page_level_data['Ship To'],
                '{{Shipment Nbr}}': str(int(float(shipment_level_data['Shipment Nbr']))),  # Remove .0 from float
                '{{PO}}': page_level_data['PO'],
                '{{DO #}}': f"{int(page_level_data['DO #']):010d}",  # Format with leading zeros (10 digits)
                '{{VAS}}': shipment_level_data['VAS'],
                '{{Original Qty}}': page_level_data['Original Qty'] + ' Units',  # Add "Units" after quantity
                '{{Label Type}}': shipment_level_data['Label Type'],
                '{{Order Type}}': shipment_level_data['Order Type'],
                '{{Pmt Term}}': shipment_level_data['Pmt Term'],
                '{{Start Ship}}': shipment_level_data['Start Ship']
            }
            
            # Create a fresh copy of template for this page
            page_doc = Document(template_path)  # type: ignore
            
            # Replace placeholders in this page
            self.replace_placeholders_in_document(page_doc, replacements)
            
            # Handle multi-page document creation
            if do_index == 1:
                # First page: use this as the main document
                main_doc = page_doc
            else:
                # Subsequent pages: add page break and copy all content with formatting
                if main_doc is not None:
                    main_doc.add_page_break()
                    self.copy_formatted_content(page_doc, main_doc)
        
        # Save document
        output_filename = f"Placard_{shipment_num}.docx"
        output_path = os.path.join(self.output_folder, output_filename)
        
        try:
            if main_doc is not None:
                main_doc.save(output_path)
                duration = (datetime.now() - start_time).total_seconds()
                self.print_with_timestamp(f"SUCCESS: Created placard document: {output_path}")
                
                # Log successful processing
                self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num,
                              do_count=total_dos, records_found=records_found,
                              status="SUCCESS", output_file=output_filename, duration=duration)
                return True
            else:
                error_msg = "No document was created"
                self.print_with_timestamp(f"ERROR: {error_msg}")
                duration = (datetime.now() - start_time).total_seconds()
                self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num,
                              do_count=total_dos, records_found=records_found,
                              status="FAILED", error_message=error_msg, duration=duration)
                return False
        except Exception as e:
            error_msg = f"ERROR saving document: {e}"
            self.print_with_timestamp(error_msg)
            duration = (datetime.now() - start_time).total_seconds()
            self.log_event("SHIPMENT_PROCESS", shipment_number=shipment_num,
                          do_count=total_dos, records_found=records_found,
                          status="FAILED", error_message=str(e), duration=duration)
            return False
    
    def get_all_unique_shipments(self) -> List[str]:
        """Get all unique shipment numbers from the dataset"""
        if self.df is None:
            return []
        
        # Get unique shipment numbers and convert to clean strings
        df_shipment_clean = self.df['Shipment Nbr'].astype(float).astype(int).astype(str)
        unique_shipments = df_shipment_clean.unique().tolist()
        
        # Filter valid shipment numbers
        valid_shipments = [s for s in unique_shipments if self.validate_shipment_number(s)]
        return sorted(valid_shipments)
    
    def process_all_shipments(self) -> Tuple[int, int]:
        """Process all shipments in the dataset"""
        self.print_with_timestamp("\n=== Processing ALL Shipments ===")
        start_time = datetime.now()
        
        all_shipments = self.get_all_unique_shipments()
        if not all_shipments:
            self.print_with_timestamp("No valid shipments found in the dataset.")
            self.log_event("BULK_PROCESS", status="FAILED", 
                          error_message="No valid shipments found", processing_mode="BULK")
            return 0, 0
        
        self.print_with_timestamp(f"Found {len(all_shipments)} unique shipments to process...")
        
        # Log bulk processing start
        self.log_event("BULK_PROCESS_START", records_found=len(all_shipments), 
                      processing_mode="BULK", status="STARTED")
        
        # Ask for confirmation
        confirm = input(f"This will generate {len(all_shipments)} placard documents. Continue? (y/n): ").strip().lower()
        if confirm not in ['y', 'yes']:
            self.print_with_timestamp("Bulk processing cancelled.")
            self.log_event("BULK_PROCESS", status="CANCELLED", 
                          error_message="User cancelled bulk processing", processing_mode="BULK")
            return 0, 0
        
        successful_count = 0
        failed_count = 0
        
        # Process each shipment
        for i, shipment_num in enumerate(all_shipments, 1):
            self.print_with_timestamp(f"\n[{i}/{len(all_shipments)}] Processing shipment: {shipment_num}")
            
            if self.process_shipment(shipment_num):
                successful_count += 1
            else:
                failed_count += 1
                
            # Show progress every 10 shipments or at the end
            if i % 10 == 0 or i == len(all_shipments):
                self.print_with_timestamp(f"Progress: {i}/{len(all_shipments)} processed ({successful_count} successful, {failed_count} failed)")
        
        # Log bulk processing completion
        duration = (datetime.now() - start_time).total_seconds()
        self.log_event("BULK_PROCESS_COMPLETE", 
                      records_found=len(all_shipments),
                      status=f"COMPLETED: {successful_count} success, {failed_count} failed",
                      processing_mode="BULK", duration=duration,
                      error_message=f"Processed {len(all_shipments)} shipments")
        
        return successful_count, failed_count
    
    def get_user_choice(self) -> str:
        """Get user choice for processing mode"""
        while True:
            try:
                self.print_with_timestamp("\nChoose an option:")
                self.print_with_timestamp("1. Enter specific shipment numbers")
                self.print_with_timestamp("2. Generate placards for ALL shipments in dataset")
                self.print_with_timestamp("3. Exit")
                
                choice = input("Enter your choice (1-3): ").strip()
                
                if choice in ['1', '2', '3']:
                    return choice
                else:
                    self.print_with_timestamp("Please enter 1, 2, or 3.")
                    
            except KeyboardInterrupt:
                self.print_with_timestamp("\nOperation cancelled by user.")
                sys.exit(0)
            except Exception as e:
                self.print_with_timestamp(f"Error reading input: {e}")
                continue
    
    def get_user_input(self) -> List[str]:
        """Get shipment numbers from user input"""
        while True:
            try:
                user_input = input("\nEnter one or more Shipment Numbers (comma-separated): ").strip()
                if not user_input:
                    self.print_with_timestamp("Please enter at least one shipment number.")
                    continue
                
                # Split by comma and clean up
                shipment_numbers = [num.strip() for num in user_input.split(',')]
                shipment_numbers = [num for num in shipment_numbers if num]  # Remove empty strings
                
                if not shipment_numbers:
                    self.print_with_timestamp("Please enter valid shipment numbers.")
                    continue
                
                return shipment_numbers
                
            except KeyboardInterrupt:
                self.print_with_timestamp("\nOperation cancelled by user.")
                sys.exit(0)
            except Exception as e:
                self.print_with_timestamp(f"Error reading input: {e}")
                continue
    
    def run(self) -> None:
        """Main execution method"""
        self.print_with_timestamp("=== Shipping Placard Generator ===")
        self.print_with_timestamp("Loading data and preparing system...")
        session_start = datetime.now()
        
        # Setup directories
        if not self.setup_directories():
            return
        
        # Initialize logging
        self.initialize_log()
        
        # Log session start
        self.log_event("SESSION_START", status="STARTED")
        
        # Load and prepare data (one-time operation)
        if not self.load_and_prepare_data():
            self.print_with_timestamp("Failed to load data. Please check the Data folder and file format.")
            self.log_event("SESSION_END", status="FAILED", 
                          error_message="Failed to load data")
            return
        
        # Check template exists
        template_path = os.path.join(self.template_folder, "placard_template.docx")
        if not os.path.exists(template_path):
            error_msg = f"Template file not found: {template_path}"
            self.print_with_timestamp(f"ERROR: {error_msg}")
            self.print_with_timestamp("Please place the template file 'placard_template.docx' in the Template folder.")
            self.log_event("SESSION_END", status="FAILED", error_message=error_msg)
            return
        
        self.print_with_timestamp("\nData loaded successfully! Ready to generate placards.")
        
        # Show dataset summary
        if self.df is not None:
            all_shipments = self.get_all_unique_shipments()
            self.print_with_timestamp(f"Dataset contains {len(all_shipments)} unique valid shipments.")
        
        # Process user requests
        total_successful = 0
        total_failed = 0
        
        while True:
            try:
                # Get user choice
                choice = self.get_user_choice()
                
                if choice == '1':
                    # Manual shipment entry
                    self.log_event("USER_CHOICE", processing_mode="MANUAL", status="SELECTED")
                    shipment_numbers = self.get_user_input()
                    
                    # Process each shipment
                    successful_count = 0
                    failed_count = 0
                    
                    for shipment_num in shipment_numbers:
                        if self.process_shipment(shipment_num):
                            successful_count += 1
                        else:
                            failed_count += 1
                    
                    total_successful += successful_count
                    total_failed += failed_count
                    
                    # Print summary
                    self.print_with_timestamp(f"\n=== Processing Summary ===")
                    self.print_with_timestamp(f"Documents created: {successful_count}")
                    self.print_with_timestamp(f"Failed inputs: {failed_count}")
                    
                    # Log manual processing summary
                    self.log_event("MANUAL_PROCESS_SUMMARY", 
                                  records_found=len(shipment_numbers),
                                  status=f"COMPLETED: {successful_count} success, {failed_count} failed",
                                  processing_mode="MANUAL")
                    
                elif choice == '2':
                    # Process all shipments
                    self.log_event("USER_CHOICE", processing_mode="BULK", status="SELECTED")
                    successful_count, failed_count = self.process_all_shipments()
                    total_successful += successful_count
                    total_failed += failed_count
                    
                    # Print summary
                    self.print_with_timestamp(f"\n=== Bulk Processing Summary ===")
                    self.print_with_timestamp(f"Documents created: {successful_count}")
                    self.print_with_timestamp(f"Failed shipments: {failed_count}")
                    
                elif choice == '3':
                    # Exit
                    self.print_with_timestamp("Exiting...")
                    break
                
                # Ask if user wants to continue (except after bulk processing or exit)
                if choice != '2':
                    continue_choice = input("\nReturn to main menu? (y/n): ").strip().lower()
                    if continue_choice not in ['y', 'yes']:
                        break
                else:
                    # After bulk processing, ask if they want to continue
                    continue_choice = input("\nReturn to main menu? (y/n): ").strip().lower()
                    if continue_choice not in ['y', 'yes']:
                        break
                    
            except KeyboardInterrupt:
                self.print_with_timestamp("\n\nOperation cancelled by user.")
                self.log_event("SESSION_END", status="INTERRUPTED", 
                              error_message="User interrupted session")
                break
            except Exception as e:
                self.print_with_timestamp(f"Unexpected error: {e}")
                self.log_event("SESSION_END", status="ERROR", error_message=str(e))
                break
        
        # Log session end
        session_duration = (datetime.now() - session_start).total_seconds()
        self.print_with_timestamp(f"\n=== Final Summary ===")
        self.print_with_timestamp(f"Total documents created: {total_successful}")
        self.print_with_timestamp(f"Total failed inputs: {total_failed}")
        self.print_with_timestamp("Thank you for using the Shipping Placard Generator!")
        
        self.log_event("SESSION_END", 
                      status="COMPLETED",
                      error_message=f"Total: {total_successful} success, {total_failed} failed",
                      duration=session_duration)


def main() -> None:
    """Main entry point"""
    generator = PlacardGenerator()
    generator.run()


if __name__ == "__main__":
    main() 